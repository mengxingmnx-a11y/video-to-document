# -*- coding: utf-8 -*-
"""Render a simple Markdown file to .docx (headings/paragraphs/lists/tables/bold).
Usage: python md_to_docx.py input.md output.docx
"""
import argparse, os, re
from docx import Document
from docx.shared import Pt

def add_runs(par, text):
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            par.add_run(part[2:-2]).bold = True
        else:
            par.add_run(part)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md")
    ap.add_argument("docx")
    a = ap.parse_args()
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(11)
    rows, lines = [], open(a.md, encoding="utf-8").read().splitlines()
    def flush():
        nonlocal rows
        if not rows:
            return
        n = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=n)
        t.style = "Light Grid Accent 1"
        for ri, r in enumerate(rows):
            for ci in range(n):
                t.cell(ri, ci).text = r[ci] if ci < len(r) else ""
        rows = []
        doc.add_paragraph()
    for line in lines:
        line = line.rstrip()
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^:?-{2,}:?$", c) for c in cells):
                continue
            rows.append(cells)
            continue
        flush()
        if line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], 3)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            add_runs(p, line[2:])
            for r in p.runs:
                r.italic = True
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
        elif line.strip() in ("", "---"):
            continue
        else:
            add_runs(doc.add_paragraph(), line)
    flush()
    doc.save(a.docx)
    print("saved:", a.docx, os.path.getsize(a.docx), "bytes")

if __name__ == "__main__":
    main()
